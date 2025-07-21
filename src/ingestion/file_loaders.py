import os
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

from src.rag import Document, Chunk


def extract_pdf_text(path: str) -> str:
    """Extract plain text from a PDF file."""
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_docx_text(path: str) -> str:
    """Extract plain text from a Word (.docx) file."""
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def extract_xlsx_text(path: str) -> str:
    """Extract cell values from an Excel (.xlsx) file as text."""
    wb = load_workbook(path, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            values = [str(cell) if cell is not None else "" for cell in row]
            lines.append("\t".join(values))
    return "\n".join(lines)


def ingest_file(path: str) -> Document:
    """Load a file and convert it into a Document for indexing."""
    extension = Path(path).suffix.lower()
    if extension == ".pdf":
        content = extract_pdf_text(path)
    elif extension == ".docx":
        content = extract_docx_text(path)
    elif extension == ".xlsx":
        content = extract_xlsx_text(path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    filename = Path(path).name
    chunk = Chunk(content=content, similarity=1.0)
    return Document(id=filename, title=filename, chunks=[chunk])
